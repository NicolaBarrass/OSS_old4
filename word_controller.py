# -*- coding: utf-8 -*-
"""
Created on Mon Mar  9 10:14:10 2020

@author: snbar
"""
from docx import Document
from docx.oxml.shared import qn
from docx.shared import Pt
from docx.shared import Inches


from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

from docxtpl import DocxTemplate

from app.classes import DateString
from app.classes import app_list

import datetime
from datetime import datetime as my_datetime


exam_meas=['hipflex','hipflexdef','hipabdflex','hipabdext','hiprotint','hiprotext','femant',
           'kneeffd','kneehyp','popliteal','kneeflexflex','kneeflexext',
           'mdfflex','mfdext','bimal']



def prep_history(appID):
    start_title='PATIENT BACKGROUND'
    end_title='SUMMARY OF FINDINGS'
    document = Document('C:\\Users\\snbar\\projects\\OSS\\Other\\ReportTemplate.docx')
    print('now printing')
    
    start=0
    stop=0
    new_paragraph=[]
    new_type=[]
    for paragraph in document.paragraphs:
#        print ('paragraph x: ', paragraph.text.lstrip())
        if paragraph.text.lstrip()==start_title.lstrip():
            start=1
#            print 'A: ', start,stop
            
        elif paragraph.text.lstrip()==end_title.lstrip():
            stop=1   
#            print 'B: ',  start,stop
        if start>0:
            start=start+1
#            print  'C: ', start,stop
            if start>2 and stop==0:
#                print paragraph.text
#                print 'stop'
#                new_paragraph.append(paragraph)
                temp_para=document.add_paragraph()
                for run in paragraph.runs:
#                    print ('text is: ', run.text)
#                    print ('bold is: ',run.bold)
#                    print ('underline is: ',run.underline)
                    temp_run=temp_para.add_run(run.text)
                    print('*',run.text)
                    temp_run.bold=run.bold
                    temp_run.underline=run.underline
                new_paragraph.append(temp_para)
                new_type.append(temp_run.underline)
#                print('#',temp_para.text)
#                print(temp_run.underline)
                

    my_text=zip(new_paragraph,new_type)  

#    for new_paragraph, new_type in my_text:
#        print(new_paragraph.text)
#        print(new_type)
        
        
#    print ('new paragraphs')
#    my_text=''
    for paragraph in new_paragraph:
        print ('#',paragraph.text)
#               
#    for t in new_type:
#        print ('##',t)
##        my_text=my_text+(paragraph.text)
##        my_text=my_text+('<br>')
        
    return my_text


def save_to_text_file(my_text):
    print ('saving to text file')
    filename="Test2.docx"
#    print filename
    document = Document()
    print (my_text)
    document.add_paragraph(my_text)
#    for paragraph in my_text:
#        p=document.add_paragraph(paragraph)
        
    document.save(filename)
    
#def prep_export(appID):
#    # get appointment for appID
#    this_a=db_get_id('Appointment',{'id':appID})#id',appID)
#    a=db_list_orderby('Appointment',(this_a.doa),{'patID':this_a.patID})
#    a_list=app_list()
#
#    for my_a in a:
#        a_list.id.append(my_a.id)
#        a_list.doa_str.append(my_a.doa_str)
#        a_list.doa.append(my_a.doa)
#        
#    return this_a,a_list



def final_report(report_var):
    doc = DocxTemplate('C:\\Users\\snbar\\projects\\OSS\\Other\\Report.docx')
    context = {  
               'years':report_var.years,
               'combined':report_var.combined2,
               'report_var':report_var
                }
    doc.render(context)
    doc.save('C:\\Users\\snbar\\projects\\OSS\\Other\\New.docx')
    